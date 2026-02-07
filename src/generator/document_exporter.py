"""
Document Exporter - Export resumes and cover letters to PDF/DOCX.

Provides clean, professional document generation with customizable formatting.
"""

import os
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentExporter:
    """
    Export documents to PDF and DOCX formats.
    
    Supports:
    - Resume export with professional formatting
    - Cover letter export
    - Clean, ATS-friendly layouts
    """
    
    def __init__(self, output_dir: str = "data/exports"):
        """
        Initialize document exporter.
        
        Args:
            output_dir: Directory for exported documents
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logging.getLogger(__name__)
    
    def export_resume_pdf(
        self,
        content: Dict[str, Any],
        filename: str = None,
    ) -> Optional[str]:
        """
        Export resume to PDF format.
        
        Args:
            content: Resume content dictionary
            filename: Output filename (auto-generated if not provided)
            
        Returns:
            Path to exported file or None if failed
        """
        if not REPORTLAB_AVAILABLE:
            self.logger.error("reportlab not installed. Install with: pip install reportlab")
            return None
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.pdf"
        
        output_path = self.output_dir / filename
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=0.5*inch,
                leftMargin=0.5*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch,
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=6,
                alignment=TA_CENTER,
            )
            
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=6,
                spaceBefore=12,
                textColor=colors.HexColor('#1a365d'),
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                alignment=TA_JUSTIFY,
            )
            
            # Header with name and contact
            if 'name' in content:
                story.append(Paragraph(content['name'], name_style))
            
            if 'contact' in content:
                contact = content['contact']
                contact_parts = []
                if contact.get('email'):
                    contact_parts.append(contact['email'])
                if contact.get('phone'):
                    contact_parts.append(contact['phone'])
                if contact.get('linkedin'):
                    contact_parts.append(contact['linkedin'])
                
                if contact_parts:
                    contact_line = " | ".join(contact_parts)
                    contact_style = ParagraphStyle(
                        'ContactStyle',
                        parent=styles['Normal'],
                        fontSize=9,
                        alignment=TA_CENTER,
                    )
                    story.append(Paragraph(contact_line, contact_style))
            
            story.append(Spacer(1, 12))
            
            # Summary
            if 'summary' in content and content['summary']:
                story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
                story.append(Paragraph(content['summary'], body_style))
            
            # Skills
            if 'skills' in content and content['skills']:
                story.append(Paragraph("SKILLS", section_style))
                skills_text = ", ".join(content['skills'][:15])
                story.append(Paragraph(skills_text, body_style))
            
            # Experience
            if 'experience' in content and content['experience']:
                story.append(Paragraph("EXPERIENCE", section_style))
                for exp in content['experience']:
                    title_line = f"<b>{exp.get('title', '')}</b> - {exp.get('company', '')}"
                    story.append(Paragraph(title_line, body_style))
                    
                    if exp.get('bullets'):
                        for bullet in exp['bullets']:
                            story.append(Paragraph(f"• {bullet}", body_style))
            
            # Education
            if 'education' in content and content['education']:
                story.append(Paragraph("EDUCATION", section_style))
                for edu in content['education']:
                    edu_line = f"<b>{edu.get('degree', '')}</b>"
                    if edu.get('field'):
                        edu_line += f" in {edu['field']}"
                    if edu.get('institution'):
                        edu_line += f" - {edu['institution']}"
                    story.append(Paragraph(edu_line, body_style))
            
            # Projects
            if 'projects' in content and content['projects']:
                story.append(Paragraph("PROJECTS", section_style))
                for proj in content['projects']:
                    proj_line = f"<b>{proj.get('name', '')}</b>"
                    if proj.get('description'):
                        proj_line += f": {proj['description']}"
                    story.append(Paragraph(proj_line, body_style))
            
            # Build PDF
            doc.build(story)
            self.logger.info(f"Resume exported to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"PDF export failed: {e}")
            return None
    
    def export_resume_docx(
        self,
        content: Dict[str, Any],
        filename: str = None,
    ) -> Optional[str]:
        """
        Export resume to DOCX format.
        
        Args:
            content: Resume content dictionary
            filename: Output filename
            
        Returns:
            Path to exported file or None if failed
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.docx"
        
        output_path = self.output_dir / filename
        
        try:
            doc = Document()
            
            # Set narrow margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Name
            if 'name' in content:
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(content['name'])
                name_run.bold = True
                name_run.font.size = Pt(18)
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact
            if 'contact' in content:
                contact = content['contact']
                contact_parts = []
                if contact.get('email'):
                    contact_parts.append(contact['email'])
                if contact.get('phone'):
                    contact_parts.append(contact['phone'])
                if contact.get('linkedin'):
                    contact_parts.append(contact['linkedin'])
                
                if contact_parts:
                    contact_para = doc.add_paragraph(" | ".join(contact_parts))
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            if 'summary' in content and content['summary']:
                doc.add_heading('PROFESSIONAL SUMMARY', level=2)
                doc.add_paragraph(content['summary'])
            
            # Skills
            if 'skills' in content and content['skills']:
                doc.add_heading('SKILLS', level=2)
                doc.add_paragraph(", ".join(content['skills'][:15]))
            
            # Experience
            if 'experience' in content and content['experience']:
                doc.add_heading('EXPERIENCE', level=2)
                for exp in content['experience']:
                    para = doc.add_paragraph()
                    title_run = para.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                    title_run.bold = True
                    
                    if exp.get('bullets'):
                        for bullet in exp['bullets']:
                            doc.add_paragraph(bullet, style='List Bullet')
            
            # Education
            if 'education' in content and content['education']:
                doc.add_heading('EDUCATION', level=2)
                for edu in content['education']:
                    edu_line = f"{edu.get('degree', '')}"
                    if edu.get('field'):
                        edu_line += f" in {edu['field']}"
                    if edu.get('institution'):
                        edu_line += f" - {edu['institution']}"
                    doc.add_paragraph(edu_line)
            
            # Projects
            if 'projects' in content and content['projects']:
                doc.add_heading('PROJECTS', level=2)
                for proj in content['projects']:
                    para = doc.add_paragraph()
                    proj_run = para.add_run(proj.get('name', ''))
                    proj_run.bold = True
                    if proj.get('description'):
                        para.add_run(f": {proj['description']}")
            
            doc.save(str(output_path))
            self.logger.info(f"Resume exported to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"DOCX export failed: {e}")
            return None
    
    def export_cover_letter_pdf(
        self,
        cover_letter: str,
        job: Dict[str, Any],
        profile: Dict[str, Any],
        filename: str = None,
    ) -> Optional[str]:
        """
        Export cover letter to PDF format.
        
        Args:
            cover_letter: Cover letter text
            job: Job data
            profile: User profile
            filename: Output filename
            
        Returns:
            Path to exported file or None if failed
        """
        if not REPORTLAB_AVAILABLE:
            self.logger.error("reportlab not installed")
            return None
        
        if not filename:
            company = job.get('company', 'company').replace(' ', '_')
            timestamp = datetime.now().strftime("%Y%m%d")
            filename = f"cover_letter_{company}_{timestamp}.pdf"
        
        output_path = self.output_dir / filename
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch,
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Header with sender info
            personal = profile.get('personal', {})
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_LEFT,
            )
            
            if personal.get('name'):
                story.append(Paragraph(personal['name'], header_style))
            if personal.get('email'):
                story.append(Paragraph(personal['email'], header_style))
            if personal.get('phone'):
                story.append(Paragraph(personal['phone'], header_style))
            
            story.append(Spacer(1, 24))
            
            # Date
            date_str = datetime.now().strftime("%B %d, %Y")
            story.append(Paragraph(date_str, header_style))
            
            story.append(Spacer(1, 24))
            
            # Recipient info
            story.append(Paragraph(f"Hiring Manager", header_style))
            story.append(Paragraph(job.get('company', ''), header_style))
            
            story.append(Spacer(1, 24))
            
            # Subject
            subject_style = ParagraphStyle(
                'SubjectStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=12,
            )
            story.append(Paragraph(
                f"<b>RE: Application for {job.get('title', 'Position')}</b>",
                subject_style
            ))
            
            # Body
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_JUSTIFY,
                spaceAfter=12,
            )
            
            # Split cover letter into paragraphs
            paragraphs = cover_letter.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
            
            # Signature
            story.append(Spacer(1, 24))
            story.append(Paragraph("Sincerely,", header_style))
            story.append(Spacer(1, 24))
            story.append(Paragraph(personal.get('name', ''), header_style))
            
            doc.build(story)
            self.logger.info(f"Cover letter exported to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Cover letter PDF export failed: {e}")
            return None
    
    def export_cover_letter_docx(
        self,
        cover_letter: str,
        job: Dict[str, Any],
        profile: Dict[str, Any],
        filename: str = None,
    ) -> Optional[str]:
        """Export cover letter to DOCX format."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not installed")
            return None
        
        if not filename:
            company = job.get('company', 'company').replace(' ', '_')
            timestamp = datetime.now().strftime("%Y%m%d")
            filename = f"cover_letter_{company}_{timestamp}.docx"
        
        output_path = self.output_dir / filename
        
        try:
            doc = Document()
            
            personal = profile.get('personal', {})
            
            # Header
            if personal.get('name'):
                doc.add_paragraph(personal['name'])
            if personal.get('email'):
                doc.add_paragraph(personal['email'])
            if personal.get('phone'):
                doc.add_paragraph(personal['phone'])
            
            doc.add_paragraph()
            
            # Date
            doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            doc.add_paragraph()
            
            # Recipient
            doc.add_paragraph("Hiring Manager")
            doc.add_paragraph(job.get('company', ''))
            doc.add_paragraph()
            
            # Subject
            subject_para = doc.add_paragraph()
            subject_run = subject_para.add_run(f"RE: Application for {job.get('title', 'Position')}")
            subject_run.bold = True
            doc.add_paragraph()
            
            # Body
            paragraphs = cover_letter.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph(personal.get('name', ''))
            
            doc.save(str(output_path))
            self.logger.info(f"Cover letter exported to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Cover letter DOCX export failed: {e}")
            return None
    
    def is_pdf_available(self) -> bool:
        """Check if PDF export is available."""
        return REPORTLAB_AVAILABLE
    
    def is_docx_available(self) -> bool:
        """Check if DOCX export is available."""
        return DOCX_AVAILABLE
