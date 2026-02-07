"""
Resume Parser - Extract structured data from PDF/DOCX resumes.

Supports PDF and DOCX formats with intelligent skill extraction,
experience parsing, and education detection.
"""

import re
import os
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

from PyPDF2 import PdfReader
from docx import Document


class ResumeParser:
    """Parse resumes and extract structured data."""
    
    # Common technical skills to look for
    TECHNICAL_SKILLS = {
        # Programming Languages
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "go",
        "rust", "scala", "kotlin", "swift", "php", "r", "matlab", "julia",
        
        # ML/AI
        "machine learning", "deep learning", "neural networks", "nlp",
        "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn",
        "opencv", "transformers", "huggingface", "llm", "gpt", "bert",
        
        # Data Science
        "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
        "jupyter", "data analysis", "data visualization", "statistics",
        "data mining", "feature engineering", "a/b testing",
        
        # Big Data
        "spark", "hadoop", "hive", "kafka", "airflow", "databricks",
        "snowflake", "redshift", "bigquery",
        
        # Databases
        "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
        "cassandra", "dynamodb", "neo4j", "sqlite",
        
        # Cloud & DevOps
        "aws", "gcp", "azure", "docker", "kubernetes", "terraform",
        "jenkins", "github actions", "ci/cd", "linux", "bash",
        
        # Web Development
        "react", "vue", "angular", "node.js", "express", "django", "flask",
        "fastapi", "spring", "html", "css", "rest api", "graphql",
        
        # Tools & Frameworks
        "git", "jira", "confluence", "agile", "scrum", "mlops",
    }
    
    # Degree patterns
    DEGREE_PATTERNS = [
        r"(?:bachelor'?s?|b\.?s\.?|b\.?tech|b\.?e\.?)\s*(?:of|in)?\s*(?:science|engineering|technology|arts)?",
        r"(?:master'?s?|m\.?s\.?|m\.?tech|m\.?e\.?)\s*(?:of|in)?\s*(?:science|engineering|technology|arts)?",
        r"(?:ph\.?d\.?|doctorate|doctor)\s*(?:of|in)?",
        r"(?:mba|m\.?b\.?a\.?)",
    ]
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume file and extract structured data.
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Dictionary with extracted resume data
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            text = self._parse_pdf(path)
        elif suffix in ['.docx', '.doc']:
            text = self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use PDF or DOCX.")
        
        # Clean the text
        text = self._clean_text(text)
        
        # Extract all components
        skills = self._extract_skills(text)
        experience = self._extract_experience(text)
        education = self._extract_education(text)
        contact = self._extract_contact(text)
        
        return {
            "file_path": str(path.absolute()),
            "file_type": suffix.lstrip('.'),
            "raw_text": text,
            "skills": skills,
            "experience": experience,
            "education": education,
            "contact": contact,
            "parsed_at": datetime.utcnow().isoformat(),
        }
    
    def _parse_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(path)
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            self.logger.error(f"Failed to parse PDF {path}: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")
    
    def _parse_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {path}: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\-\@\+\(\)\/]', '', text)
        return text.strip()
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.TECHNICAL_SKILLS:
            # Use word boundaries for accurate matching
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill.title())
        
        # Also look for skill variations
        skill_variations = {
            "python": ["python3", "python 3"],
            "machine learning": ["ml", "machinelearning"],
            "deep learning": ["dl", "deeplearning"],
            "natural language processing": ["nlp"],
            "computer vision": ["cv"],
            "javascript": ["js", "es6", "ecmascript"],
            "typescript": ["ts"],
            "node.js": ["nodejs", "node js"],
            "react": ["react.js", "reactjs"],
            "vue": ["vue.js", "vuejs"],
            "postgresql": ["postgres"],
            "mongodb": ["mongo"],
            "kubernetes": ["k8s"],
            "continuous integration": ["ci/cd", "cicd"],
        }
        
        for main_skill, variations in skill_variations.items():
            for var in variations:
                if var in text_lower and main_skill.title() not in found_skills:
                    found_skills.append(main_skill.title())
                    break
        
        return sorted(list(set(found_skills)))
    
    def _extract_experience(self, text: str) -> Dict[str, Any]:
        """Extract work experience information."""
        experience = {
            "years": 0,
            "positions": [],
            "companies": [],
        }
        
        # Try to find years of experience patterns
        year_patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
            r'(?:experience|exp)(?:ience)?:?\s*(\d+)\+?\s*(?:years?|yrs?)',
        ]
        
        text_lower = text.lower()
        for pattern in year_patterns:
            match = re.search(pattern, text_lower)
            if match:
                experience["years"] = int(match.group(1))
                break
        
        # Extract date ranges to estimate experience
        date_pattern = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}\s*[-–]\s*(?:present|current|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4})'
        date_matches = re.findall(date_pattern, text_lower)
        
        if date_matches and experience["years"] == 0:
            # Estimate from date ranges
            total_months = 0
            for date_range in date_matches:
                months = self._estimate_months_from_range(date_range)
                total_months += months
            experience["years"] = round(total_months / 12, 1)
        
        return experience
    
    def _estimate_months_from_range(self, date_range: str) -> int:
        """Estimate months from a date range string."""
        month_map = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }
        
        try:
            parts = re.split(r'[-–]', date_range)
            if len(parts) != 2:
                return 0
            
            start_match = re.search(r'([a-z]+)[a-z]*\.?\s*(\d{4})', parts[0])
            end_match = re.search(r'(?:present|current)|([a-z]+)[a-z]*\.?\s*(\d{4})', parts[1])
            
            if not start_match:
                return 0
            
            start_month = month_map.get(start_match.group(1)[:3], 1)
            start_year = int(start_match.group(2))
            
            if 'present' in parts[1] or 'current' in parts[1]:
                end_month = datetime.now().month
                end_year = datetime.now().year
            elif end_match and end_match.group(1):
                end_month = month_map.get(end_match.group(1)[:3], 12)
                end_year = int(end_match.group(2))
            else:
                return 0
            
            months = (end_year - start_year) * 12 + (end_month - start_month)
            return max(0, months)
        except Exception:
            return 0
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information."""
        education = []
        text_lower = text.lower()
        
        # Look for degree patterns
        for pattern in self.DEGREE_PATTERNS:
            matches = re.finditer(pattern, text_lower)
            for match in matches:
                degree_text = match.group(0)
                
                # Try to find associated field of study
                context = text_lower[max(0, match.start()-50):min(len(text_lower), match.end()+100)]
                
                field = None
                field_patterns = [
                    r'(?:computer science|cs)',
                    r'(?:data science)',
                    r'(?:machine learning)',
                    r'(?:artificial intelligence|ai)',
                    r'(?:software engineering)',
                    r'(?:information technology|it)',
                    r'(?:electrical engineering)',
                    r'(?:mathematics|math)',
                    r'(?:statistics)',
                ]
                
                for fp in field_patterns:
                    if re.search(fp, context):
                        field = re.search(fp, context).group(0)
                        break
                
                # Extract year
                year_match = re.search(r'20\d{2}', context)
                year = int(year_match.group(0)) if year_match else None
                
                education.append({
                    "degree": degree_text.strip(),
                    "field": field,
                    "year": year,
                })
        
        return education
    
    def _extract_contact(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information."""
        contact = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
        }
        
        # Email pattern
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact["email"] = email_match.group(0).lower()
        
        # Phone pattern (various formats)
        phone_patterns = [
            r'[\+]?[\d\s\-\(\)]{10,}',
            r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                phone = re.sub(r'[^\d\+]', '', match.group(0))
                if len(phone) >= 10:
                    contact["phone"] = phone
                    break
        
        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w\-]+', text.lower())
        if linkedin_match:
            contact["linkedin"] = f"https://{linkedin_match.group(0)}"
        
        # GitHub
        github_match = re.search(r'github\.com/[\w\-]+', text.lower())
        if github_match:
            contact["github"] = f"https://{github_match.group(0)}"
        
        return contact
    
    def get_skill_categories(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills into groups."""
        categories = {
            "programming": [],
            "ml_ai": [],
            "data_science": [],
            "databases": [],
            "cloud_devops": [],
            "web_development": [],
            "tools": [],
        }
        
        skill_to_category = {
            "programming": ["python", "java", "javascript", "typescript", "c++", "c#", 
                           "ruby", "go", "rust", "scala", "kotlin", "swift", "php", "r"],
            "ml_ai": ["machine learning", "deep learning", "tensorflow", "pytorch", 
                     "keras", "scikit-learn", "nlp", "computer vision", "neural networks"],
            "data_science": ["pandas", "numpy", "scipy", "matplotlib", "data analysis",
                            "statistics", "jupyter", "data visualization"],
            "databases": ["sql", "mysql", "postgresql", "mongodb", "redis", 
                         "elasticsearch", "cassandra", "dynamodb"],
            "cloud_devops": ["aws", "gcp", "azure", "docker", "kubernetes", 
                            "terraform", "jenkins", "ci/cd", "linux"],
            "web_development": ["react", "vue", "angular", "node.js", "django", 
                               "flask", "fastapi", "html", "css", "rest api"],
            "tools": ["git", "jira", "confluence", "agile", "scrum"],
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            categorized = False
            
            for category, cat_skills in skill_to_category.items():
                if skill_lower in cat_skills:
                    categories[category].append(skill)
                    categorized = True
                    break
            
            if not categorized:
                categories["tools"].append(skill)
        
        return {k: v for k, v in categories.items() if v}


# Convenience function
def parse_resume(file_path: str) -> Dict[str, Any]:
    """Parse a resume file and return structured data."""
    parser = ResumeParser()
    return parser.parse(file_path)
